import json
import os
import re
import sys
import difflib
from argparse import ArgumentParser, ArgumentTypeError
import traceback
from docx import Document, table
from collections import namedtuple
import openpyxl
from openpyxl.workbook import Workbook

from openpyxl.cell import Cell, MergedCell


Metadata_ids = namedtuple("Metadata_ids", "sections, data_elements, countries, combos")


def get_country_id(country: str, countries_ids: dict):
    country_key = None
    for key in countries_ids:
        if country in key:
            country_key = key

    if country_key:
        return countries_ids[country_key]
    else:
        error = f'Can\'t find orgUnit id for country: {country}'
        raise ValueError(error)


def get_data_element_id(de: str, data_elements_ids: dict):
    if de in data_elements_ids:
        return data_elements_ids[de]
    else:
        print(f'Can\'t find id for dataElement: {de}')
        print(f'Closest candidates: {difflib.get_close_matches(de, data_elements_ids)}')
        return None


def get_metadata_ids(workbook: Workbook):
    sections_id_dict = {}
    data_elements_id_dict = {}
    countries_id_dict = {}
    combos_id_dict = {}
    sheet = workbook['Metadata']

    for row in sheet.iter_rows(min_row=2, min_col=1, max_col=5, values_only=True):
        identifier = row[0]
        type_col = row[1]
        name = cleanup_string(row[2])

        Option_set = row[4] if row[4] else False

        if type_col == 'sections':
            sections_id_dict[name] = identifier

        if type_col == 'categoryOptionCombos':
            combos_id_dict[name] = identifier

        if type_col == 'dataElements' and not Option_set:
            data_elements_id_dict[name] = identifier

        if type_col == 'organisationUnit':
            countries_id_dict[name] = identifier

    return Metadata_ids(sections_id_dict, data_elements_id_dict, countries_id_dict, combos_id_dict)


def make_matched_values(tables_data: list[dict], coverage_tables_data: dict, ids: Metadata_ids):
    country_id = get_country_id(COUNTRY, ids.countries)
    default_combo_id = ids.combos['default']

    data = {}
    data[country_id] = {}
    data[country_id][YEAR] = {}

    for table_data in tables_data:
        for de_name, value in table_data.items():
            if not de_name in ids.sections:
                de_id = get_data_element_id(de_name, ids.data_elements)
                if not de_id:
                    pass
                if de_id not in data[country_id][YEAR]:
                    data[country_id][YEAR][de_id] = {}

            data[country_id][YEAR][de_id][default_combo_id] = value

    for coverage_year, coverage_data in coverage_tables_data.items():
        if coverage_year not in data[country_id]:
            data[country_id][coverage_year] = {}

        for de_name, value in coverage_data.items():
            if not de_name in ids.sections:
                de_id = get_data_element_id(de_name, ids.data_elements)
                if not de_id:
                    pass
                if de_id not in data[country_id][coverage_year]:
                    data[country_id][coverage_year][de_id] = {}

            data[country_id][coverage_year][de_id][default_combo_id] = value

    return data


def write_org_unit(last_cell: Cell, matched_values: dict):
    for country_id, country_data in matched_values.items():
        for year in country_data:
            new_cell = last_cell.offset(row=1, column=0)
            new_cell.value = f'=_{country_id}'

            last_cell = new_cell


def write_years(last_cell: Cell, matched_values: dict):
    org_unit_years_row = {}

    for country_id, country_data in matched_values.items():
        if country_id not in org_unit_years_row:
            org_unit_years_row[country_id] = {}

        for year in country_data:
            new_cell = last_cell.offset(row=1, column=0)
            new_cell.value = year

            last_cell = new_cell

            if year not in org_unit_years_row[country_id]:
                org_unit_years_row[country_id][year] = last_cell.row

    return org_unit_years_row


def write_data(col_indicator: str, col_combo: str, col_first_value_cell: Cell, matched_values: dict, org_unit_years_row: dict):

    for country_id, country_data in matched_values.items():
        for year, data_elements in country_data.items():
            for indicator_id, indicator_combos in data_elements.items():
                if indicator_id == col_indicator:
                    for combo_id, value in indicator_combos.items():
                        ids = combo_id.split('|') if '|' in combo_id else combo_id
                        if col_combo in ids or (col_combo == 'Xr12mI7VPn3' and combo_id == 'gEWtgad4feW'):
                            row_offset = org_unit_years_row[country_id][year]-5
                            new_cell = col_first_value_cell.offset(row=row_offset, column=0)
                            new_cell.value = value

    return col_first_value_cell


def write_values(workbook: Workbook, matched_values: dict):
    sheet = workbook['Data Entry']
    workbook.active = workbook['Data Entry']

    for index, col in enumerate(sheet.iter_cols(min_row=4)):
        if index == 0:
            last_cell = col[-1]
            write_org_unit(last_cell, matched_values)

        if index == 1:
            last_cell = col[-1]
            org_unit_years_row = write_years(last_cell, matched_values)

        if index == 2:
            pass
        if index > 2:
            if not isinstance(col[0], MergedCell):
                col_indicator = str(col[0].value).split('=_')[-1]
            col_combo = str(col[1].value).split('=_')[-1]
            col_first_value_cell = col[-1]

            write_data(col_indicator, col_combo, col_first_value_cell, matched_values, org_unit_years_row)

    workbook.save(OUT_FILENAME)


def get_country_and_year(document: Document):
    """
    Reads the first line from the source DOCX file and returns the metadata country and year.

    :param document: The document loaded from the source DOCX file.
    :type document: Document
    :return: The country and year of the data in the DOCX file.
    :rtype: tuple (str, str)
    """

    try:
        first_paragraph = document.paragraphs[0]
        first_line = first_paragraph.text

        country, year = first_line.replace(")", "").split(" (")

        if not year.isnumeric() or len(year) != 4:
            raise ValueError(f'Invalid year format for: {year}')

        return country, year
    except Exception as e:
        print('First line of the doc should be: Country (year)', file=sys.stderr)
        traceback.print_exc()
        sys.exit(1)


def fix_references_format(references: str):
    """
    Adds " to the third element of each references line.

    :param references: The string containing the references row values.
    :type references: str
    :return: The references string with formatting applied.
    :rtype: str
    """

    new_references = ""

    for reference in references.splitlines():
        reference_partition = re.split(r': *(<.*?>) *', reference, 1)

        if not reference_partition[-1].startswith('\"'):
            reference_partition[-1] = '"' + reference_partition[-1]

        if not reference_partition[-1].endswith('\"'):
            reference_partition[-1] += '"'

        new_references += f'{reference_partition[0]}: {reference_partition[1]} {reference_partition[2]}\n'

    return new_references


def extract_longtext_tables(document: Document):
    """
    Extracts tables from the source DOCX file and returns them as a list of dictionaries.
    Ignores fields starting with "Internal".

    :param document: The document loaded from the source DOCX file.
    :type document: Document
    :return: A list of dictionaries representing the tables in the .docx file.
    :rtype: list
    """

    tables_data_list = []

    for table in document.tables:
        if len(table.columns) != 2:
            continue

        table_data = {}
        for row in table.rows:
            key = row.cells[0].text.rstrip()

            if key.startswith(INTERNAL):
                continue

            value = row.cells[1].text.rstrip()

            if key == "References":
                value = fix_references_format(value)

            if key and value:
                table_data[key] = value
            else:
                debug(
                    f'Empty row with {"DE: " + key if key else ""}{" and " if key and value else ""}{"value: " + value if value else ""} in source file')

        tables_data_list.append(table_data)

    return tables_data_list


def table_is_target(table: table, test_header: list):
    debug("table_is_target header len: ", len(table.columns), len(test_header))
    if len(table.columns) != len(test_header):
        return False

    table_header = [cell.text.strip() for cell in table.rows[0].cells]
    debug("table_is_target headers: ", table_header, test_header)
    if table_header == test_header:
        return True


def get_charges_in_coverage_table_data(document: Document, header_dict: dict, target_headers: list):
    table_data = {}
    year_count = {}

    for table in document.tables:
        header_keys = list(header_dict.keys())
        header_data_elements = list(header_dict.values())

        if table_is_target(table, header_keys):
            for row_id, row in enumerate(table.rows):
                for cell_id, cell in enumerate(row.cells):
                    # Skip header
                    if row_id == 0:
                        continue

                    if cell_id == 0:
                        year = cell.text.strip()
                        if year not in table_data:
                            table_data[year] = {}
                            year_count[year] = 1
                        else:
                            # NOTE: Limited to COVERAGE_TABLE_MAX items as its the limit of the data entry form
                            if year_count[year] >= COVERAGE_TABLE_MAX:
                                error(
                                    f"Ccharges in coverage table has too many entries for year {year} (max {COVERAGE_TABLE_MAX} entries for year)"
                                )
                                error(f"Discarded row:")
                                error(f"{' | '.join([cleanup_string(cell.text) for cell in row.cells])}")
                                continue
                            else:
                                year_count[year] += 1

                    if cell_id in target_headers:
                        data_element = f'{header_data_elements[cell_id]} ({year_count[year]})'
                        table_data[year][data_element] = cell.text.strip()
            return table_data


def cleanup_string(text: str):
    return ' '.join(str(text).split())


def extract_user_charges_by_type_table(document: Document):
    header_list = [
        "Type of health care",
        "User charges apply",
        "Type of user charge",
        "Reduced user charges",
        "Exemptions from user charges",
        "Cap on user charges",
    ]

    type_of_heath_care_dict = {
        "Outpatient primary care visits (text)": "(Primary Care)",
        "Outpatient specialist visits (text)": "(Specialist Visits)",
        "Outpatient emergency visits (text)": "(Emergency Visits)",
        "Outpatient prescribed medicines (text)": "(Outpatient Medicines)",
        "Medical products (text)": "(Medical Products)",
        "Diagnostic tests (cat)": "(Diagnostics Tests)",
        "Dental care visits (text)": "(Dental Visits)",
        "Dental care treatment (text)": "(Dental Care Treatment)",
        "Inpatient care (text)": "(Inpatient Care)",
        "Inpatient medicines (text)": "(Inpatient Medicines)",
    }

    table_data = {}

    for table in document.tables:

        if table_is_target(table, header_list):
            for row_id, row in enumerate(table.rows):
                if row_id == 0:
                    continue
                for cell_id, cell in enumerate(row.cells):
                    if cell_id == 0:
                        type_of_heath_care = type_of_heath_care_dict[cleanup_string(cell.text)]
                        continue
                    data_element = f'{header_list[cell_id]} {type_of_heath_care}'
                    table_data[data_element] = cell.text.strip()
            return table_data


def add_to_coverage_tables_data(coverage_tables_data: dict, new_data: dict):
    if new_data:
        for year, values in new_data.items():
            coverage_tables_data[year] = values

    return dict(sorted(coverage_tables_data.items(), reverse=True))


def extract_charges_in_coverage_upto19_table(document: Document):
    changes_in_coverage_2019_header = {
        "Year": "",
        "Month": "",
        "Coverage policy area": "Area of change in coverage policy pre2019",
        "Policy change": "Policy of change in coverage policy pre2019",
        "Health services targeted": "Health services targeted in change in coverage policy pre2019",
        "People targeted": "People targeted in change in coverage policy pre2019",
        "Coverage policy area (cat)": "",
        "Health services targeted (cat)": "",
        "People targeted (cat)": ""
    }

    target_headers = [2, 3, 4, 5]
    table_data = get_charges_in_coverage_table_data(document, changes_in_coverage_2019_header, target_headers)

    if table_data:
        return table_data
    else:
        error('Cant find "Changes in coverage policy up to and including 2019" table.')
        return None


def extract_charges_in_coverage_since20_table(document: Document):
    changes_in_coverage_2020_header = {
        "Year": "",
        "Month": "",
        "Coverage policy area": "Area of change in coverage policy",
        "Policy change": "Policy of change in coverage policy",
        "Health services targeted": "Health services targeted in change in coverage policy",
        "People targeted": "People targeted in change in coverage policy",
        "Was this a response to the COVID-19 pandemic?": "Was this a response to the COVID-19 pandemic?",
        "Coverage policy area (cat)": "",
        "Health services targeted (cat)": "",
        "People targeted (cat)": "",
        "Was this a response to the COVID-19 pandemic? (cat)": ""
    }

    target_headers = [2, 3, 4, 5, 6]
    table_data = get_charges_in_coverage_table_data(document, changes_in_coverage_2020_header, target_headers)

    if table_data:
        return table_data
    else:
        error('Cant find "Changes in coverage policy since 2020" table.')
        return None


def debug(*msg):
    if DEBUG:
        with open(LOG_FILE, "a") as log_file:
            print(*msg, file=log_file)


def error(*msg):
    print(*msg, file=sys.stderr)


def dump_json_var(var):
    return json.dumps(var, indent=2)


def filepath_exists(filepath):
    return os.path.isfile(filepath)


def get_template_path(parser: ArgumentParser, xlsx_template: str):
    if not xlsx_template:
        if filepath_exists(DEFAULT_TEMPLATE):
            xlsx_template = DEFAULT_TEMPLATE
        else:
            parser.error(f'The default template: {DEFAULT_TEMPLATE} doesn\'t exist')
    elif not filepath_exists(xlsx_template):
        parser.error(f'The template: {xlsx_template} doesn\'t exist')

    return xlsx_template


def get_coverage_max(parser: ArgumentParser, value: int):
    if value:
        if value <= 0:
            parser.error(f'The coverage_max value must be positive.')
        else:
            return value
    else:
        return 10


def main():
    """
    Parses command-line arguments and extracts tables from a .docx file.
    """
    parser = ArgumentParser(description='Process DOCX files into "Bulk Load" XLSX files. \
                                     The script needs a template, it either can be supplied with the --xlsx_template \
                                     argument or by placing a template named "Qualitative_Data_UHCPW_Template.xlsx" \
                                     in the same folder as the script.\
                                     Outputs to a XLSX file named <COUNTRY>_<YEAR>_Qualitative_Data.xlsx.')
    parser.add_argument('docx_filename', type=str,
                        help='The path to the DOCX source file.')
    parser.add_argument('-x', '--xlsx_template', type=str,
                        help='Bulk Load Qualitative XLSX template file path, \
                            if empty tries to open "Qualitative_Data_UHCPW_Template.xlsx"')
    parser.add_argument('-d', '--debug', action='store_true',
                        help='Print debug logs into a "log.json" file.')
    parser.add_argument('-c', '--coverage_max', type=int,
                        help='Number of coverage policy table entries per year, by default 10, must be positive.')
    args = parser.parse_args()

    if not filepath_exists(args.docx_filename):
        parser.error(f'The source file: {args.docx_filename} doesn\'t exist')

    global OUT_FILENAME, DEFAULT_TEMPLATE, DEBUG, LOG_FILE, COUNTRY, YEAR, INTERNAL, COVERAGE_TABLE_MAX
    INTERNAL = 'Internal'
    DEFAULT_TEMPLATE = 'Qualitative_Data_UHCPW_Template.xlsx'
    DEBUG = args.debug

    if DEBUG:
        LOG_FILE = "log.json"
        f = open(LOG_FILE, 'w')
        f.close()

    COVERAGE_TABLE_MAX = get_coverage_max(parser, args.coverage_max)

    args.xlsx_template = get_template_path(parser, args.xlsx_template)

    debug('Source file:', args.docx_filename)
    debug('Template file:', args.xlsx_template)

    document = Document(args.docx_filename)

    COUNTRY, YEAR = get_country_and_year(document)
    OUT_FILENAME = f'{COUNTRY}_{YEAR}_Qualitative_Data.xlsx'

    debug('Output file:', OUT_FILENAME)
    debug('Country:', COUNTRY, 'Year:', YEAR)

    longtext_tables_data = extract_longtext_tables(document)
    debug('longtext_tables_data:\n', dump_json_var(longtext_tables_data))

    tables_data = longtext_tables_data
    user_charges_by_type_data = extract_user_charges_by_type_table(document)
    debug('user_charges_by_type_data:\n', dump_json_var(user_charges_by_type_data))
    if user_charges_by_type_data:
        tables_data.append(user_charges_by_type_data)

    coverage_tables_data = {}
    charges_in_coverage_upto19_data = extract_charges_in_coverage_upto19_table(document)
    coverage_tables_data = add_to_coverage_tables_data(coverage_tables_data, charges_in_coverage_upto19_data)

    charges_in_coverage_since20_data = extract_charges_in_coverage_since20_table(document)
    coverage_tables_data = add_to_coverage_tables_data(coverage_tables_data, charges_in_coverage_since20_data)

    debug('coverage_tables_data:\n', dump_json_var(coverage_tables_data))

    try:
        wb = openpyxl.load_workbook(args.xlsx_template)
    except Exception as e:
        debug("openpyxl exception: ", e)
        traceback.print_exc()
        sys.exit(1)

    ids = get_metadata_ids(wb)

    debug(f'sections ids:\n len: {len(ids.sections)}\n values:\n', dump_json_var(ids.sections))
    debug(f'data_elements ids:\n len: {len(ids.data_elements)}\n values:\n', dump_json_var(ids.data_elements))
    debug(f'countries ids:\n len: {len(ids.countries)}\n values:\n', dump_json_var(ids.countries))
    debug(f'combos ids:\n len: {len(ids.combos)}\n values:\n', dump_json_var(ids.combos))

    debug('tables_data:\n', dump_json_var(tables_data))

    matched_values = make_matched_values(tables_data, coverage_tables_data, ids)

    debug(f'matched_values:\n', dump_json_var(matched_values))

    write_values(wb, matched_values)


if __name__ == '__main__':
    main()
